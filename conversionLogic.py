from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import docx
import fitz
import io
import logging
import subprocess
import tempfile
import os
from pdf2docx import Converter
from logger_config import configure_logger

logger = configure_logger(__name__)

libreoffice_path = '/Applications/LibreOffice.app/Contents/MacOS/soffice'


def txt_to_pdf(input_data, output_stream):
    """
    Converts plain text to a PDF file.

    This function takes plain text data and an output stream (a BytesIO object),
    and writes the text to the output stream as a formatted PDF file.

    Parameters:
    - input_data (str): The text to be converted into PDF format.
    - output_stream (io.BytesIO): The output stream where the PDF content will be written.

    Returns:
    None: The function writes the output directly to the provided output stream.
    """

    c = canvas.Canvas(output_stream, pagesize=letter)
    width, height = letter
    styles = getSampleStyleSheet()
    styleN = styles['Normal']

    lines = input_data.split('\n')

    for i, line in enumerate(lines):
        y = height - (i * 20) - 20  # Increase line spacing to 20 units
        text = Paragraph(line.strip(), styleN)
        text.wrapOn(c, width - 20, height)
        text.drawOn(c, 10, y)

    c.save()


def txt_to_docx(input_data):
    """
    Converts plain text to a DOCX document with basic formatting.

    Parameters:
    - input_data (str): The text to be converted into DOCX format,
                        assuming the first line is the title.

    Returns:
    A DOCX document object containing the input text with formatting.
    """
    doc = docx.Document()
    lines = input_data.split('\n')

    if lines:
        doc.add_heading(lines[0], level=1)

        for line in lines[1:]:
            paragraph = doc.add_paragraph(line)


    return doc


def docx_to_txt(input_stream):
    """
    Converts a DOCX document to plain text.

    This function reads a DOCX file from the given BytesIO object, extracts all text from
    the document, and returns it as a single plain text string.

    Parameters:
    - input_stream (io.BytesIO): A BytesIO object containing the DOCX document to be converted.

    Returns:
    str: A string containing the text extracted from the DOCX document.
    """
    # Ensure the stream is at the start
    input_stream.seek(0)

    # Load the DOCX file from the BytesIO object
    doc = docx.Document(input_stream)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


def pdf_to_txt(input_stream):
    """
    Extracts all text from a PDF file.

    Parameters:
    - input_stream (io.BytesIO): The stream of the PDF document to be converted.

    Returns:
    str: A string containing all text extracted from the PDF document.
    """

    text = ''
    with fitz.open(stream=input_stream) as doc:
        for page in doc:
            text += page.get_text()

    return text


def docx_to_pdf(input_stream, output_stream):
    """
    Converts a DOCX document to a PDF file using LibreOffice's command-line interface.

    This function takes a DOCX document from a BytesIO object, converts it to PDF format using
    LibreOffice, and writes the result to another BytesIO object specified as the output stream.

    Parameters:
    - input_stream (io.BytesIO): The input stream containing the DOCX document to be converted.
    - output_stream (io.BytesIO): The output stream where the converted PDF content will be written.

    Raises:
    - RuntimeError: If the conversion process fails or if the expected PDF file is not created.
    """

    with tempfile.TemporaryDirectory() as tmpdirname:
        docx_path = os.path.join(tmpdirname, "input.docx")
        pdf_path = os.path.join(tmpdirname, "input.pdf")

        with open(docx_path, 'wb') as tmp_docx:
            tmp_docx.write(input_stream.read())

        try:
            subprocess.run([libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', tmpdirname, docx_path],
                           check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e.stderr}")
            raise RuntimeError("Failed to convert DOCX to PDF using LibreOffice.")

        if not os.path.exists(pdf_path):
            logger.error("Expected PDF file was not created.")
            raise FileNotFoundError("Expected PDF file was not created.")

        with open(pdf_path, 'rb') as pdf_file:
            output_stream.write(pdf_file.read())


def pdf_to_docx(input_stream, output_stream):
    """
    Converts a PDF file to DOCX format using the pdf2docx library.

    Parameters:
    - input_stream (io.BytesIO): The input stream containing the PDF document to be converted.
    - output_stream (io.BytesIO): The output stream where the converted DOCX content will be written.

    Raises:
    - RuntimeError: If the conversion process fails.
    """
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf_file, \
            tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_docx_file:

        temp_pdf_file.write(input_stream.read())
        input_pdf_path = temp_pdf_file.name
        output_docx_path = temp_docx_file.name

        try:
            cv = Converter(input_pdf_path)
            cv.convert(output_docx_path)
            cv.close()
        except Exception as e:
            logger.error(f"PDF to DOCX conversion failed: {e}")
            raise RuntimeError("Failed to convert PDF to DOCX.")

        with open(output_docx_path, 'rb') as docx_file:
            output_stream.write(docx_file.read())

        # Cleanup
        os.unlink(input_pdf_path)
        os.unlink(output_docx_path)


def txt_to_rtf(input_stream, output_stream):
    """
    Converts a plain text input stream to RTF format and writes to an output stream.
8
    Parameters:
    - input_stream (io.BytesIO): The input stream containing the plain text to be converted.
    - output_stream (io.BytesIO): The output stream where the RTF content will be written.
    """
    # RTF header, document area start, and footer. Adjust formatting as needed.
    rtf_header = r"{\rtf1\ansi\deff0{\fonttbl{\f0\fswiss Arial;}}\f0\pard "
    rtf_footer = r"\par}"

    input_text = input_stream.getvalue().decode('utf-8')

    # Escape RTF special characters in the input text
    escaped_text = input_text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}').replace('\n', '\\par\n')

    # Assemble the RTF content
    rtf_content = f"{rtf_header}{escaped_text}{rtf_footer}"

    output_stream.write(rtf_content.encode('utf-8'))


def docx_to_rtf(input_stream):
    """
    Converts DOCX content from an input stream to RTF format using Pandoc,
    and returns the RTF content in an output stream.

    Parameters:
    - input_stream (io.BytesIO): Input stream containing the DOCX file content.

    Returns:
    io.BytesIO: Output stream containing the RTF formatted content, or None in case of failure.
    """
    try:
        # Use a temporary directory to hold the files
        with tempfile.TemporaryDirectory() as tmpdir:
            input_tmp_path = os.path.join(tmpdir, 'document.docx')
            output_tmp_path = os.path.join(tmpdir, 'document.rtf')

            # Write the input stream content to a temporary file
            with open(input_tmp_path, 'wb') as input_tmp:
                input_tmp.write(input_stream.getvalue())

            # Construct and execute the Pandoc command
            command = ['pandoc', '-s', input_tmp_path, '-o', output_tmp_path]
            subprocess.run(command, check=True)

            # Read the converted RTF content back into an output stream
            output_stream = io.BytesIO()
            with open(output_tmp_path, 'rb') as output_file:
                output_stream.write(output_file.read())

            output_stream.seek(0)
            return output_stream

    except subprocess.CalledProcessError:
        raise ("Conversion failed: Pandoc encountered an error while converting the document.")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")
    return None


def pdf_to_rtf(input_stream):
    doc = fitz.open(stream=input_stream, filetype="pdf")
    rtf_text = "{\\rtf1\\ansi\\deff0\\pard "

    for page in doc:
        text = page.get_text()
        rtf_text += text.replace('\n', '\\par\n') + '\\par\n'

    rtf_text += "}"
    output_stream = io.BytesIO(rtf_text.encode('utf-8'))
    return output_stream


def rtf_to_docx(input_stream, output_stream):
    """
    Converts an RTF document to a DOCX file using LibreOffice's command-line interface.

    This function takes an RTF document from a BytesIO object, converts it to DOCX format using
    LibreOffice, and writes the result to another BytesIO object specified as the output stream.

    Parameters:
    - input_stream (io.BytesIO): The input stream containing the RTF document to be converted.
    - output_stream (io.BytesIO): The output stream where the converted DOCX content will be written.

    Raises:
    - RuntimeError: If the conversion process fails.
    """

    with tempfile.TemporaryDirectory() as tmpdirname:
        rtf_path = os.path.join(tmpdirname, "input.rtf")
        docx_path = os.path.join(tmpdirname, "input.docx")

        logger.info(f"Temporary RTF path: {rtf_path}")
        logger.info(f"Temporary DOCX path: {docx_path}")

        with open(rtf_path, 'wb') as tmp_rtf:
            tmp_rtf.write(input_stream.read())
        logging.info(f"Written RTF content to temporary file")
        cmd = [libreoffice_path, '--headless', '--convert-to', 'docx', '--outdir', tmpdirname, rtf_path]

        logger.info(f"Executing command: {' '.join(cmd)}")

        try:
            subprocess.run(cmd, check=True, capture_output=True)
            logger.info("Conversion command executed successfully")
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e.output}")
            raise RuntimeError("LibreOffice conversion failed.")

        if not os.path.exists(docx_path):
            logger.error("The expected DOCX file was not created.")
            raise FileNotFoundError("The expected DOCX file was not created.")

        with open(docx_path, 'rb') as docx_file:
            output_stream.write(docx_file.read())

        return output_stream


def rtf_to_pdf(input_stream, output_stream):
    with tempfile.TemporaryDirectory() as tmpdirname:
        rtf_path = os.path.join(tmpdirname, "input.rtf")
        pdf_path = os.path.join(tmpdirname, "input.pdf")

        logger.info(f"Temporary RTF path: {rtf_path}")
        logger.info(f"Temporary PDF path: {pdf_path}")

        with open(rtf_path, 'wb') as tmp_rtf:
            tmp_rtf.write(input_stream.read())
        logging.info(f"Written RTF content to temporary file")
        cmd = [libreoffice_path, '--headless', '--convert-to', 'pdf', '--outdir', tmpdirname, rtf_path]

        logger.info(f"Executing command: {' '.join(cmd)}")

        try:
            subprocess.run(cmd, check=True, capture_output=True)
            logger.info("Conversion command executed successfully")
        except subprocess.CalledProcessError as e:
            logger.error(f"LibreOffice conversion failed: {e.output}")
            raise RuntimeError("LibreOffice conversion failed.")

        if not os.path.exists(pdf_path):
            logger.error("The expected DOCX file was not created.")
            raise FileNotFoundError("The expected DOCX file was not created.")

        with open(pdf_path, 'rb') as pdf_file:
            output_stream.write(pdf_file.read())

        return output_stream
