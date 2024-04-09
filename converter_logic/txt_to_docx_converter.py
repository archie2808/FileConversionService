from base_converter import BaseConverter
from io import BytesIO
import docx


class TXTToDocxConverter(BaseConverter):
    def convert(self, output_stream: BytesIO):
        doc = docx.Document()
        input_text = self.input_stream.getvalue().decode('utf-8')
        lines = input_text.split('\n')

        if lines:
            doc.add_heading(lines[0], level=1)

            for line in lines[1:]:
                doc.add_paragraph(line)

        # Save DOCX to an output stream
        doc.save(output_stream)
