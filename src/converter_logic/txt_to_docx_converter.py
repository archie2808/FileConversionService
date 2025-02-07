from io import BytesIO
import docx
from ..base_converter import BaseConverter


class TXTToDocxConverter(BaseConverter):
    """
    Converts plain text to a DOCX document.

    This class uses the python-docx library to create a DOCX document from the provided
    plain text input stream.
    """

    def convert(self, output_stream: BytesIO):
        """
        Reads plain text from the input stream, creates a DOCX document, and saves it to the output stream.

        Parameters:
            output_stream (BytesIO): The stream where the DOCX file will be saved..
        """
        doc = docx.Document()
        input_text = self.input_stream.getvalue().decode('utf-8')
        lines = input_text.split('\n')

        if lines:
            doc.add_heading(lines[0], level=1)

            for line in lines[1:]:
                doc.add_paragraph(line)


        doc.save(output_stream)
